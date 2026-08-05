"""
DOCX extraction + chunking ด้วย Docling
==========================================
แทนที่ extraction.py + hierarchical_chunking.py เดิม (เฉพาะไฟล์ .docx เท่านั้น
ไฟล์ .pdf ยังใช้ pypdf เดิม ไม่เปลี่ยน เพราะ Docling ต้องพึ่ง backend หนัก
(torch, docling-ibm-models, OCR) เฉพาะตอนอ่าน PDF ที่ไม่มีโครงสร้าง XML
ให้อ่านตรงๆ เหมือน DOCX ต้องใช้ vision model ช่วยเข้าใจ layout)

ทำไมเลือกทางนี้ (ติดตั้งแบบเบา ไม่ใช่ `pip install docling` ตรงๆ):
`pip install docling` ธรรมดาจะดึง torch + torchvision + transformers +
CUDA libraries เต็มชุด (~3-5GB) เพราะ default ไปรวม backend สำหรับ PDF/OCR
ทั้งที่เราต้องการแค่ DOCX ถ้าจะย้ายมาใช้จริง ให้ติดตั้งแบบนี้แทน:

    pip install "docling-slim[format-docx,feat-chunking,convert-core]" pypdfium2

(pypdfium2 ต้องเพิ่มเองแยกต่างหาก - extras ของ docling-slim ไม่ได้ประกาศ
ไว้ให้ แต่ backend DOCX เรียกใช้จริงตอนจัดการรูปภาพ/drawing ที่ฝังในไฟล์
เจอจาก error จริงตอน deploy ไม่ใช่แค่ทดสอบเฉยๆ - ไฟล์เล็กมาก ~4MB
ไม่กระทบเรื่องความเบาของการติดตั้งเลย)

ยืนยันแล้วว่าไม่ต้องมี torch เลยก็ทำงานได้ (import ได้เป็น `docling`
เหมือนแพ็กเกจเต็ม แค่ชื่อบน PyPI ต่างกัน)

ใช้ MsWordDocumentBackend ตรงๆ (ไม่ใช้ DocumentConverter ทั่วไป) เพราะ
DocumentConverter import backend ของ PDF (docling_parse) แบบ eager
ตั้งแต่ตอน import module ถึงจะไม่ได้ใช้ PDF เลยก็ตาม ทำให้ error
ModuleNotFoundError ทันทีถ้าไม่ได้ติดตั้ง PDF backend ไว้

สำคัญ: heal_heading_styles() ต้องเรียกก่อนเสมอ - เจอจริงจากไฟล์ผู้ใช้ว่า
paragraph.style เป็น None ได้ (style reference resolve ไม่ได้) กรณีนี้
Docling ไม่ crash แต่ก็ "ไม่นับเป็นหัวข้อ" ให้เงียบๆ เช่นกัน (ต่างจาก
extraction.py เดิมที่มี fallback pattern ช่วยจับ) เลยต้อง "รักษา" ไฟล์
ก่อนส่งให้ Docling อ่าน โดยบังคับ apply Heading style ให้ paragraph ที่
ข้อความ match pattern "ปีที่ X...ภาคการศึกษาที่ Y" แต่ style ยังไม่ใช่
Heading จริง (ครอบคลุมทั้งเคส style=None และเคส style=Normal ธรรมดา)
"""

import io
import re

from docx import Document as DocxDocument
from docling.backend.msword_backend import MsWordDocumentBackend
from docling.datamodel.base_models import InputFormat
from docling.datamodel.document import InputDocument
from docling_core.transforms.chunker.hierarchical_chunker import HierarchicalChunker


# pattern เดียวกับที่ extraction.py เดิมใช้เป็น fallback - เก็บไว้ที่นี่
# เพราะตอนนี้ทำหน้าที่เป็น "ตัวรักษาไฟล์ก่อนส่งเข้า Docling" แทน
FALLBACK_HEADING_PATTERN = re.compile(r"ปีที่\s*\d+.*ภาคการศึกษาที่\s*\d+")


def heal_heading_styles(file_bytes: bytes) -> tuple[bytes, int]:
    """
    บังคับ apply Heading 1 ให้ paragraph ที่ข้อความ match pattern หัวข้อ
    เทอม/ปี แต่ style ปัจจุบันไม่ใช่ Heading จริง (ครอบคลุมทั้ง style=None
    ที่เจอจริงจากไฟล์ และ style='Normal' ทั่วไป) คืนค่า (bytes ที่แก้แล้ว,
    จำนวน paragraph ที่ถูกแก้) - จำนวนที่คืนมามีไว้ log/debug ดูว่า
    healing ทำงานจริงกี่จุดในแต่ละเอกสาร
    """
    doc = DocxDocument(io.BytesIO(file_bytes))
    fixed_count = 0

    for paragraph in doc.paragraphs:
        style = paragraph.style
        is_heading_style = (
            style is not None and style.name.lower().startswith("heading")
        )
        if not is_heading_style and FALLBACK_HEADING_PATTERN.search(paragraph.text):
            paragraph.style = doc.styles["Heading 1"]
            fixed_count += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), fixed_count


def _chunk_to_lines(chunk) -> list[str]:
    """
    แปลง 1 raw chunk (จาก HierarchicalChunker) เป็น list ของบรรทัด
    (1 บรรทัด = 1 record/รายวิชา) - ต้องแยก 2 กรณีเพราะ doc_items มี
    ได้ 2 ชนิดที่โครงสร้างไม่เหมือนกันเลย (เจอจริงจากการทดสอบ 2 ไฟล์
    ที่โครงสร้างเอกสารต่างกัน):

    - TextItem (มี .text ตรงๆ): เอกสารที่แต่ละรายวิชาเป็น paragraph
      เดี่ยวๆ แยกกัน (มักถูกตัดเป็นหลาย TextItem ต่อ 1 รายวิชา เพราะ
      bold/italic ทำให้ Word แยก run) -> รวม TextItem ที่ติดกันเป็น
      1 บรรทัดเดียว (เพราะเป็นรายวิชาเดียวกัน)
    - TableItem (ไม่มี .text แต่มี .data.table_cells): เอกสารที่เป็น
      ตารางหลายคอลัมน์ (1 แถว = หลายรายวิชาคนละ cell) -> แยกแต่ละ cell
      เป็นคนละบรรทัด (เพราะเป็นคนละรายวิชากัน ห้ามรวมกัน)
    """
    lines: list[str] = []
    text_parts: list[str] = []

    def flush_text_parts():
        if text_parts:
            lines.append(" ".join(text_parts))
            text_parts.clear()

    for item in chunk.meta.doc_items:
        text = getattr(item, "text", None)
        if text is not None:
            if text.strip():
                text_parts.append(text.strip())
            continue

        table_data = getattr(item, "data", None)
        if table_data is not None and hasattr(table_data, "table_cells"):
            flush_text_parts()
            for cell in table_data.table_cells:
                if cell.text and cell.text.strip():
                    lines.append(cell.text.strip())

    flush_text_parts()
    return lines


def _merge_chunks_by_heading(chunks, max_chars: int = 1500) -> list[dict]:
    """
    รวม chunk ที่ HierarchicalChunker ให้มา (1 chunk ต่อ 1 element ที่ตรวจเจอ
    เช่น 1 paragraph ต่อ 1 รายวิชา) เข้าด้วยกัน ถ้าอยู่ใต้ heading เดียวกัน
    (ใช้ c.meta.headings ตรงๆ ไม่ต้อง parse string เดา)

    ทำไมต้องมีขั้นนี้: เจอจริงจากไฟล์ผู้ใช้ว่าบางเอกสาร แต่ละรายวิชาเป็น
    คนละ element แยกกัน (คนละ paragraph) ทำให้ HierarchicalChunker เปล่าๆ
    สร้าง 1 chunk ต่อ 1 รายวิชา - เทอมที่มี 9 วิชา กลายเป็น 9 chunk แยก
    ไม่มีตัวไหนมีบริบท "เทอมนี้มีวิชาอะไรบ้าง" ครบเลย พอ retrieve มาแค่
    k=3 เลยได้คำตอบไม่ครบ (ตอบวิชาเดียวแทนที่จะเป็นทั้งเทอม)

    สร้างข้อความจาก doc_item.text ดิบผ่าน _chunk_to_lines() (ไม่ใช่
    chunk.text) เพราะ chunk.text ที่ Docling สร้างให้มี bold/italic ของ
    Word แปลงเป็น markdown **/* ติดมาด้วย (เจอจริง: "**308-101**" หลุด
    เข้าไปในคำตอบ LLM ตรงๆ) doc_item.text/cell.text เป็นข้อความดิบ
    สะอาด ไม่มี markdown เลย แก้ปัญหานี้ไปพร้อมกันในตัว

    ห้ามผสมหลาย record (รายวิชา) เข้าบรรทัดเดียวกัน - 1 บรรทัด = 1
    รายวิชาเสมอ ตัดกลุ่มใหม่เมื่อ heading เปลี่ยนหรือรวมแล้วจะเกิน
    max_chars (ไม่ตัดกลางบรรทัดเด็ดขาด)
    """
    groups: list[tuple[tuple, list[str]]] = []
    current_heading: tuple | None = None
    current_lines: list[str] = []

    for chunk in chunks:
        heading = tuple(chunk.meta.headings) if chunk.meta.headings else ()
        if heading != current_heading:
            if current_lines:
                groups.append((current_heading, current_lines))
            current_heading = heading
            current_lines = []
        current_lines.extend(_chunk_to_lines(chunk))

    if current_lines:
        groups.append((current_heading, current_lines))

    results: list[dict] = []
    for heading, lines in groups:
        heading_str = " > ".join(heading) if heading else ""
        batch: list[str] = []
        batch_len = 0

        def flush_batch():
            if not batch:
                return
            body = "\n".join(batch)
            parent = f"{heading_str}\n{body}" if heading_str else body
            results.append({"chunk_text": body, "parent_text": parent})

        for line in lines:
            if batch and batch_len + len(line) > max_chars:
                flush_batch()
                batch, batch_len = [], 0
            batch.append(line)
            batch_len += len(line) + 1

        flush_batch()

    return results


def process_docx(file_bytes: bytes, document_name: str) -> list[dict]:
    """
    รับ .docx bytes -> คืน list[dict] พร้อม insert ลง document_chunk โดยตรง
    แต่ละ dict มี:
        chunk_text  : เนื้อหาดิบของ chunk นี้ (ไม่มี heading นำหน้า)
        parent_text : เนื้อหา + heading context นำหน้า (สิ่งที่ LLM จะอ่าน
                      และสิ่งที่เอาไป embed ด้วย)

    v2: เพิ่ม _merge_chunks_by_heading() ต่อจาก HierarchicalChunker แทนที่
    จะใช้ผลจาก chunker ตรงๆ - ดู docstring ของฟังก์ชันนั้นสำหรับเหตุผล
    """
    healed_bytes, _ = heal_heading_styles(file_bytes)

    stream = io.BytesIO(healed_bytes)
    in_doc = InputDocument(
        path_or_stream=stream,
        format=InputFormat.DOCX,
        backend=MsWordDocumentBackend,
        filename=document_name,
    )
    backend = MsWordDocumentBackend(in_doc, path_or_stream=stream)
    dl_doc = backend.convert()

    chunker = HierarchicalChunker()
    raw_chunks = list(chunker.chunk(dl_doc))

    return _merge_chunks_by_heading(raw_chunks)