"""
Text extraction จากไฟล์ PDF และ Word (.docx)

v2: เปลี่ยนจาก "chunking แบบ regex เดาโครงสร้าง" มาเป็น
"header-based chunking" โดยใช้ Heading Style จริงที่ผู้ใช้เลือกไว้ใน Word
(Heading 1, Heading 2, ... ที่เห็นในแถบ Styles ของ Microsoft Word)

ทำไมเปลี่ยน: ของเดิมต้องเขียน regex ดักจับหัวข้อเอง (เช่น "ปีที่ \d+")
ใช้ได้เฉพาะเอกสารรูปแบบตายตัว พอเจอเอกสารแบบอื่นต้องเขียน regex ใหม่ทุกครั้ง
แต่ python-docx อ่าน paragraph.style.name ได้ตรงๆ ว่าย่อหน้านั้นเป็น Heading ระดับไหน
(ผู้ใช้เลือกเองตอนพิมพ์ Word อยู่แล้ว) แม่นยำกว่าการเดาจาก pattern ข้อความมาก

รูปแบบข้อมูลที่คืนออกไปตอนนี้เปลี่ยนจาก "str เดียวยาวๆ" เป็น
list[tuple[int | None, str]] คือ list ของ (heading_level, ข้อความ)
- heading_level = None        -> เป็น "เนื้อหาปกติ" ไม่ใช่หัวข้อ
- heading_level = 1, 2, 3...  -> เป็นหัวข้อระดับนั้น (มาจาก Heading 1/2/3 ใน Word)

ไฟล์ PDF ไม่มีข้อมูล style ติดมาด้วย (pypdf อ่านได้แค่ตัวอักษรล้วน)
เลยส่งกลับเป็น heading_level=None ทุกบรรทัด (เหมือนเดิม ไม่มี header ให้จับ)

ต้องลง dependency เพิ่ม: pip install pypdf python-docx
"""

import io
import re

from pypdf import PdfReader
from docx import Document as DocxDocument


class UnsupportedFileTypeError(Exception):
    pass


# จับชื่อ style แบบ "Heading 1", "Heading 2", ...
# (Word ตั้งชื่อ style ภายในเป็นอังกฤษเสมอ แม้ UI จะเป็นเวอร์ชันภาษาไทยก็ตาม)
_HEADING_STYLE_RE = re.compile(r"^heading\s*(\d+)$", re.IGNORECASE)


def _detect_heading_level(style_name: str | None) -> int | None:
    """
    แปลงชื่อ style ของย่อหน้า -> heading level (int) หรือ None ถ้าไม่ใช่ heading
    """
    name = (style_name or "").strip()

    match = _HEADING_STYLE_RE.match(name)
    if match:
        return int(match.group(1))

    # "Title" ใน Word คือหัวเรื่องใหญ่สุดของเอกสาร ถือเป็น heading ระดับบนสุด
    # (เทียบเท่า Heading 1 เพื่อให้ path ไม่ขาดตอน)
    if name.lower() == "title":
        return 1

    return None


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """แปลง PDF (bytes) เป็นข้อความล้วน รวมทุกหน้า (PDF ไม่มี heading style ให้อ่าน)"""
    reader = PdfReader(io.BytesIO(file_bytes))
    pages_text = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages_text.append(page_text)
    return "\n".join(pages_text)


def extract_structured_from_docx(file_bytes: bytes) -> list[tuple[int | None, str]]:
    """
    แปลง Word (.docx) เป็น list ของ (heading_level, ข้อความ) ทีละย่อหน้า
    โดยอ่าน Heading Style จริงที่ผู้ใช้เลือกไว้ใน Word (ไม่เดาจาก regex อีกต่อไป)
    """
    doc = DocxDocument(io.BytesIO(file_bytes))

    result: list[tuple[int | None, str]] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style is not None else None
        level = _detect_heading_level(style_name)

        result.append((level, text))

    # ตาราง: ยังไม่มีแนวคิด "heading" ในตาราง จึงใส่เป็นเนื้อหาปกติ (level=None) เหมือนเดิม
    # ต่อท้ายหลังย่อหน้าทั้งหมด (python-docx ไม่ได้ให้ลำดับ paragraph/table ปนกันตรงๆ ง่ายๆ)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                result.append((None, row_text))

    return result


def extract_text(filename: str, file_bytes: bytes) -> list[tuple[int | None, str]]:
    """
    เลือกวิธี extract ตามนามสกุลไฟล์ รองรับ .pdf, .docx
    คืนค่าเป็น list[(heading_level, text)] เสมอ ไม่ว่าไฟล์ต้นทางจะเป็นชนิดไหน
    เพื่อให้ ingestion pipeline ใช้ chunker ตัวเดียวกันได้ทั้งระบบ
    """
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext == "pdf":
        raw_text = extract_text_from_pdf(file_bytes)
        # PDF ไม่มี heading style -> ทุกบรรทัดเป็น level=None (เนื้อหาปกติ)
        return [(None, line) for line in raw_text.split("\n") if line.strip()]

    elif ext == "docx":
        return extract_structured_from_docx(file_bytes)

    else:
        raise UnsupportedFileTypeError(
            f"ไม่รองรับไฟล์นามสกุล .{ext} (รองรับเฉพาะ .pdf, .docx)"
        )